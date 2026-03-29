"""Document Source Extractor - PDF/DOCX/TXT with multi-agent extraction.

Поддержка:
- PDF (текстовые через pypdf, сканированные — TODO: GigaChat vision)
- DOCX (текст + таблицы через python-docx)
- TXT (plain text с auto-detect encoding)

См. docs/SOURCE_NODE_CONCEPT.md — раздел "📄 4. Document Dialog".
"""
import io
import logging
import re
import time
import zipfile
from pathlib import Path
from typing import Any

from app.sources.base import (
    BaseSource,
    ExtractionResult,
    ValidationResult,
    TableData,
)

logger = logging.getLogger(__name__)

# MIME type → document_type mapping
MIME_TO_DOCTYPE: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "text/plain": "txt",
}

# Extension → document_type mapping
EXT_TO_DOCTYPE: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "txt",
    ".text": "txt",
    ".md": "txt",
    ".csv": "txt",
}


def detect_document_type(config: dict[str, Any]) -> str:
    """Detect document type from config (mime_type, filename, or explicit document_type)."""
    # Explicit type
    if config.get("document_type"):
        return config["document_type"]
    
    # From MIME type
    mime = config.get("mime_type", "")
    if mime in MIME_TO_DOCTYPE:
        return MIME_TO_DOCTYPE[mime]
    
    # From filename extension
    filename = config.get("filename", "")
    for ext, dtype in EXT_TO_DOCTYPE.items():
        if filename.lower().endswith(ext):
            return dtype
    
    return "txt"  # fallback


# Sentinel: ZIP — это Excel, не Word (диалог «Документ» такие файлы не обрабатывает)
_DOC_TYPE_EXCEL_XLSX = "xlsx_wrong_dialog"

# Скан заголовка в поиске сигнатуры (часть PDF идёт с префиксом до %PDF)
_SNIFF_HEAD = 65536


def _as_bytes_for_extraction(file_content: bytes | memoryview | Path) -> bytes:
    """Единый байтовый буфер для парсеров (в т.ч. local storage → Path). Извлечение без temp-файлов."""
    if isinstance(file_content, bytes):
        return file_content
    if isinstance(file_content, memoryview):
        return file_content.tobytes()
    if isinstance(file_content, Path):
        return file_content.read_bytes()
    raise TypeError(f"Unsupported file_content type: {type(file_content)}")


def _pdf_magic_offset(data: bytes) -> int:
    """0 — если PDF с начала; иначе позиция %PDF в первых _SNIFF_HEAD байтах, либо -1."""
    if len(data) >= 4 and data[:4] == b"%PDF":
        return 0
    head = data[:_SNIFF_HEAD] if len(data) > _SNIFF_HEAD else data
    return head.find(b"%PDF")


def _trim_pdf_stream(data: bytes) -> bytes:
    """Убрать мусор до начала PDF (некоторые экспорты добавляют префикс)."""
    off = _pdf_magic_offset(data)
    if off > 0:
        return data[off:]
    return data


def detect_document_type_from_content(config: dict[str, Any], file_content: bytes) -> str:
    """Определить тип по сигнатуре байтов, затем по MIME/имени.

    Браузеры и прокси часто шлют неверный Content-Type; имя файла тоже бывает без расширения.
    Без этого PDF ошибочно идёт в ветку TXT/DOCX и даёт 400.
    """
    if _pdf_magic_offset(file_content) >= 0:
        return "pdf"
    if len(file_content) >= 4 and file_content[:2] == b"PK":
        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                names = zf.namelist()
            if "word/document.xml" in names:
                return "docx"
            if "xl/workbook.xml" in names:
                return _DOC_TYPE_EXCEL_XLSX
        except zipfile.BadZipFile:
            pass
    return detect_document_type(config)


class DocumentSource(BaseSource):
    """Document file source handler (PDF/DOCX/TXT)."""
    
    source_type = "document"
    display_name = "Документ"
    icon = "📄"
    description = "PDF, DOCX, TXT с AI-извлечением"
    
    async def validate_config(self, config: dict[str, Any]) -> ValidationResult:
        """Validate document source config."""
        errors = []
        
        if not config.get("file_id"):
            errors.append("Необходимо указать file_id")
        
        document_type = detect_document_type(config)
        if document_type not in ("pdf", "docx", "txt"):
            errors.append(f"Неподдерживаемый тип документа: {document_type}")
        
        if errors:
            return ValidationResult.failure(errors)
        return ValidationResult.success()
    
    async def extract(
        self,
        config: dict[str, Any],
        file_content: bytes | memoryview | Path | None = None,
        **kwargs
    ) -> ExtractionResult:
        """Extract text and tables from document.
        
        Args:
            config: Source config with file_id, filename, document_type, etc.
            file_content: Сырые байты (или Path для local storage). Без temp-файлов — только BytesIO в памяти.
            
        Returns:
            ExtractionResult with text and tables.
        """
        start_time = time.time()
        
        if file_content is None:
            return ExtractionResult.failure("Содержимое файла не предоставлено")

        try:
            file_content = _as_bytes_for_extraction(file_content)
        except TypeError as e:
            return ExtractionResult.failure(str(e))

        if len(file_content) == 0:
            return ExtractionResult.failure("Файл пустой")

        try:
            document_type = detect_document_type_from_content(config, file_content)
            filename = config.get("filename", "document")

            if document_type == _DOC_TYPE_EXCEL_XLSX:
                return ExtractionResult.failure(
                    "Файл — Excel (xlsx). В диалоге «Документ» поддерживаются PDF, DOCX и TXT; "
                    "для таблиц Excel используйте источник «Excel»."
                )

            logger.info(f"📄 Extracting document: {filename} (type={document_type}, size={len(file_content)} bytes)")
            
            if document_type == "pdf":
                text, tables = await self._extract_pdf(file_content, config)
            elif document_type == "docx":
                text, tables = await self._extract_docx(file_content, config)
            else:
                text, tables = self._extract_txt(file_content, config)
            
            total_rows = sum(len(t.rows) for t in tables)
            summary_line = f"Документ «{filename}»"
            if tables:
                summary_line += f". Найдено таблиц: {len(tables)}"
            if total_rows:
                summary_line += f", всего строк: {total_rows}"
            summary_line += f". Длина текста: {len(text)} символов."
            
            extraction_time = int((time.time() - start_time) * 1000)
            
            logger.info(f"✅ Document extracted: {len(text)} chars, {len(tables)} tables, {extraction_time}ms")
            
            return ExtractionResult(
                success=True,
                text=text,
                tables=tables,
                extraction_time_ms=extraction_time,
                metadata={
                    "document_type": document_type,
                    "filename": filename,
                    "summary": summary_line,
                    "text_length": len(text),
                    "table_count": len(tables),
                    "total_rows": total_rows,
                    "page_count": config.get("_page_count", None),
                    "is_scanned": bool(config.get("_is_scanned", False)),
                }
            )
            
        except Exception as e:
            logger.exception(f"❌ Document extraction failed: {e}")
            return ExtractionResult.failure(f"Ошибка извлечения из документа: {str(e)}")
    
    # ─── PDF ───────────────────────────────────────────────────────────────
    
    async def _extract_pdf(
        self, content: bytes, config: dict[str, Any]
    ) -> tuple[str, list[TableData]]:
        """Extract text and tables from PDF using pypdf (только память, без temp-файлов)."""
        from pypdf import PdfReader

        content = _trim_pdf_stream(content)
        reader = PdfReader(io.BytesIO(content), strict=False)
        # Не прерываем извлечение, если decrypt("") не вернул успех: часть PDF с owner-only
        # или нестандартным шифрованием даёт NOT_DECRYPTED в одной версии pypdf, но текст
        # всё равно читается через extract_text (локально и в Docker должно совпадать).
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception as e:
                logger.warning("PDF decrypt('') raised: %s — continuing with extract_text", e)

        total_pages = len(reader.pages)
        config["_page_count"] = total_pages
        
        # Parse page range if specified
        page_range = config.get("page_range")
        pages_to_extract = self._parse_page_range(page_range, total_pages)
        
        text_parts: list[str] = []
        for page_num in pages_to_extract:
            page = reader.pages[page_num]
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
        
        text = "\n\n".join(text_parts)
        
        # Try extracting tables from text using heuristic
        tables = self._extract_tables_from_text(text)
        
        config["_is_scanned"] = False

        # If no text at all, document might be scanned
        if not text.strip():
            text = "[Документ не содержит текстового слоя. Возможно, это скан — требуется OCR.]"
            logger.warning("PDF has no text layer, likely scanned")
            config["_is_scanned"] = True
        
        return text, tables
    
    # ─── DOCX ──────────────────────────────────────────────────────────────
    
    async def _extract_docx(
        self, content: bytes, config: dict[str, Any]
    ) -> tuple[str, list[TableData]]:
        """Extract text and tables from DOCX using python-docx."""
        # DOCX — ZIP (OOXML). Старый бинарный .doc часто ошибочно помечается как docx по расширению.
        if len(content) < 4 or not content.startswith(b"PK\x03\x04"):
            raise ValueError(
                "Файл не является DOCX (Office Open XML). Старый формат .doc не поддерживается — "
                "откройте в Word и сохраните как DOCX."
            )

        from docx import Document

        doc = Document(io.BytesIO(content))
        
        # Extract paragraphs (preserving structure)
        text_parts: list[str] = []
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                # Mark headings
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    text_parts.append(f"## {stripped}")
                else:
                    text_parts.append(stripped)
        
        text = "\n\n".join(text_parts)
        
        # Extract tables from document
        tables: list[TableData] = []
        for i, table in enumerate(doc.tables):
            parsed = self._parse_docx_table(table, index=i)
            if parsed:
                tables.append(parsed)
        
        return text, tables
    
    def _parse_docx_table(self, table, index: int) -> TableData | None:
        """Parse a single DOCX table into TableData."""
        all_rows: list[list[str]] = []
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            all_rows.append(cell_texts)
        
        if len(all_rows) < 1:
            return None
        
        # First row as headers
        header_row = all_rows[0]
        
        # Deduplicate header names (DOCX merged cells can produce duplicates)
        seen: dict[str, int] = {}
        clean_headers: list[str] = []
        for h in header_row:
            name = h if h else f"Колонка"
            if name in seen:
                seen[name] += 1
                name = f"{name}_{seen[name]}"
            else:
                seen[name] = 0
            clean_headers.append(name)
        
        columns = [{"name": h, "type": "text"} for h in clean_headers]
        
        rows: list[dict[str, Any]] = []
        for row_data in all_rows[1:]:
            row_dict = {}
            for k, val in enumerate(row_data):
                if k < len(clean_headers):
                    row_dict[clean_headers[k]] = val
            rows.append(row_dict)
        
        if not columns:
            return None
        
        # Try to detect column types from data
        for col in columns:
            col_name = col["name"]
            values = [r.get(col_name, "") for r in rows if r.get(col_name)]
            if values:
                col["type"] = self._detect_column_type(values)
        
        return TableData(
            id=f"table_{index + 1}",
            name=f"Таблица {index + 1}",
            columns=columns,
            rows=rows,
        )
    
    # ─── TXT ───────────────────────────────────────────────────────────────
    
    def _extract_txt(
        self, content: bytes, config: dict[str, Any]
    ) -> tuple[str, list[TableData]]:
        """Extract text from plain text file with encoding detection."""
        # Try common encodings
        for encoding in ["utf-8", "windows-1251", "cp1251", "latin1"]:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = content.decode("utf-8", errors="replace")
        
        tables = self._extract_tables_from_text(text)
        return text, tables
    
    # ─── Helpers ───────────────────────────────────────────────────────────
    
    @staticmethod
    def _parse_page_range(page_range: str | None, total_pages: int) -> list[int]:
        """Parse page range string like '1-5' or '1,3,5-7' into 0-based page indices."""
        if not page_range:
            return list(range(total_pages))
        
        pages: set[int] = set()
        for part in page_range.split(","):
            part = part.strip()
            if "-" in part:
                start_str, end_str = part.split("-", 1)
                start = max(1, int(start_str.strip()))
                end = min(total_pages, int(end_str.strip()))
                pages.update(range(start - 1, end))
            else:
                p = int(part)
                if 1 <= p <= total_pages:
                    pages.add(p - 1)
        
        return sorted(pages) if pages else list(range(total_pages))
    
    @staticmethod
    def _detect_column_type(values: list[str]) -> str:
        """Detect column type from sample values."""
        numeric_count = 0
        date_count = 0
        
        for v in values[:20]:  # sample first 20
            v = v.strip()
            if not v:
                continue
            # Number check
            try:
                v_clean = v.replace(" ", "").replace(",", ".").replace("%", "").replace("₽", "").replace("$", "")
                float(v_clean)
                numeric_count += 1
                continue
            except ValueError:
                pass
            # Date check
            if re.match(r"^\d{1,4}[.\-/]\d{1,2}[.\-/]\d{1,4}$", v):
                date_count += 1
        
        total = len([v for v in values[:20] if v.strip()])
        if total == 0:
            return "text"
        if numeric_count / total > 0.7:
            return "number"
        if date_count / total > 0.5:
            return "date"
        return "text"
    
    @staticmethod
    def _extract_tables_from_text(text: str) -> list[TableData]:
        """Heuristic: try to find table-like structures in plain text.
        
        Looks for lines with consistent separators (|, tab, multiple spaces).
        """
        tables: list[TableData] = []
        lines = text.split("\n")
        
        # Look for pipe-delimited tables
        table_lines: list[str] = []
        for line in lines:
            stripped = line.strip()
            if "|" in stripped and stripped.count("|") >= 2:
                # Skip separator lines like |---|---|
                if re.match(r"^[\|\-\s:+]+$", stripped):
                    continue
                table_lines.append(stripped)
            else:
                if len(table_lines) >= 2:
                    # Process accumulated table
                    parsed = DocumentSource._parse_text_table(table_lines, len(tables))
                    if parsed:
                        tables.append(parsed)
                table_lines = []
        
        # Don't forget last table
        if len(table_lines) >= 2:
            parsed = DocumentSource._parse_text_table(table_lines, len(tables))
            if parsed:
                tables.append(parsed)
        
        return tables
    
    @staticmethod
    def _parse_text_table(lines: list[str], index: int) -> TableData | None:
        """Parse pipe-delimited text lines into TableData."""
        parsed_rows: list[list[str]] = []
        for line in lines:
            cells = [c.strip() for c in line.split("|")]
            # Remove empty first/last from leading/trailing pipes
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            if cells:
                parsed_rows.append(cells)
        
        if len(parsed_rows) < 2:
            return None
        
        headers = parsed_rows[0]
        columns = [{"name": h or f"Колонка {i+1}", "type": "text"} for i, h in enumerate(headers)]
        
        rows: list[dict[str, Any]] = []
        for row_data in parsed_rows[1:]:
            row_dict: dict[str, Any] = {}
            for k, val in enumerate(row_data):
                if k < len(headers):
                    row_dict[headers[k] or f"Колонка {k+1}"] = val
            rows.append(row_dict)
        
        return TableData(
            id=f"table_{index + 1}",
            name=f"Таблица {index + 1}",
            columns=columns,
            rows=rows,
        )
    
    def get_dialog_schema(self) -> dict[str, Any]:
        """Get JSON Schema for document dialog."""
        return {
            "type": "object",
            "properties": {
                "file": {
                    "type": "file",
                    "accept": ".pdf,.docx,.txt",
                    "title": "Документ",
                },
                "extraction_prompt": {
                    "type": "string",
                    "title": "Что извлечь",
                    "description": "Опишите, какие данные нужно извлечь из документа",
                },
            },
        }
